#!/usr/bin/env python3
"""
Per-document-part redaction coverage tests.

Confirms that DocxMap._build()/apply_replacements() correctly scan AND
write back track changes across every DOCX part that can carry running
text: the main body, table cells (including cells inside headers/footers/
footnotes), headers, footers, footnotes, endnotes, textboxes, and content
controls (both block-level and inline SDTs).

Review comments (word/comments.xml paragraph content) are a documented
exception: comment text is never scanned for PII by the redaction
pipeline. By default this is masked because MetadataCleaningSettings
removes the whole comments part; if a caller explicitly retains comments
(e.g. --no-clean-review-comments-visible), a REVIEW_COMMENTS_NOT_SCANNED
warning must appear in the audit report so the gap is disclosed rather
than leaking PII silently. See docs/USER_GUIDE.md, "Redaction coverage by
document part," for the documented matrix this test enforces.
"""
import json
import os
import tempfile
import unittest
import zipfile
from xml.etree import ElementTree as ET

try:
    from docx import Document
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

try:
    from lxml import etree as LET
    LXML_AVAILABLE = True
except Exception:
    LXML_AVAILABLE = False

try:
    from marcut import pipeline
    IMPORTS_SUCCESS = True
except Exception:
    IMPORTS_SUCCESS = False


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
V_NS = "urn:schemas-microsoft-com:vml"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"


def _qn(ns: str, tag: str) -> str:
    return f"{{{ns}}}{tag}"


# Each marker uses a distinct local-part so a substring search can't
# accidentally match a different part's marker.
MARKERS = {
    "body": "bodymark@example.com",
    "table": "tablemark@example.com",
    "header": "headermark@example.com",
    "footer": "footermark@example.com",
    "footnote": "footnotemark@example.com",
    "endnote": "endnotemark@example.com",
    "comment": "commentmark@example.com",
    "textbox": "textboxmark@example.com",
    "sdt": "sdtmark@example.com",
    "sdt_inline": "sdtinlinemark@example.com",
}


@unittest.skipUnless(
    DOCX_AVAILABLE and LXML_AVAILABLE and IMPORTS_SUCCESS,
    "python-docx, lxml, or marcut not available",
)
class TestDocxPartRedactionCoverage(unittest.TestCase):
    """Build one DOCX with a distinct, known PII marker in every part, run
    the rules-only pipeline, and assert track changes land in every part
    except the documented review-comments exception."""

    def setUp(self):
        tmpdir = tempfile.TemporaryDirectory()
        self.addCleanup(tmpdir.cleanup)
        self.work_dir = tmpdir.name
        self.input_docx = os.path.join(self.work_dir, "input.docx")
        self._build_probe_docx(self.input_docx)

    # ---- fixture construction ---------------------------------------

    @staticmethod
    def _read_zip_entry(path: str, name: str) -> bytes:
        with zipfile.ZipFile(path) as zf:
            return zf.read(name)

    @staticmethod
    def _patch_zip(path: str, updates: dict, new_entries: dict | None = None) -> None:
        new_entries = new_entries or {}
        tmp = path + ".tmp"
        with zipfile.ZipFile(path, "r") as zin:
            existing = {i.filename for i in zin.infolist()}
            with zipfile.ZipFile(tmp, "w", compression=zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    data = updates.get(item.filename, zin.read(item.filename))
                    zout.writestr(item, data)
                for name, data in new_entries.items():
                    if name not in existing:
                        zout.writestr(name, data)
        os.replace(tmp, path)

    @staticmethod
    def _add_content_type_override(content_types_xml: bytes, part_name: str, content_type: str) -> bytes:
        root = ET.fromstring(content_types_xml)
        for override in root.findall(_qn(CT_NS, "Override")):
            if override.get("PartName") == part_name:
                return content_types_xml
        ET.SubElement(root, _qn(CT_NS, "Override"), {
            "PartName": part_name,
            "ContentType": content_type,
        })
        return ET.tostring(root, encoding="utf-8", xml_declaration=True)

    @staticmethod
    def _add_relationship(rels_xml: bytes, rel_id: str, rel_type: str, target: str) -> bytes:
        root = ET.fromstring(rels_xml)
        for rel in root.findall(_qn(REL_NS, "Relationship")):
            if rel.get("Id") == rel_id:
                return rels_xml
        ET.SubElement(root, _qn(REL_NS, "Relationship"), {
            "Id": rel_id,
            "Type": rel_type,
            "Target": target,
        })
        return ET.tostring(root, encoding="utf-8", xml_declaration=True)

    def _build_probe_docx(self, path: str) -> None:
        doc = Document()
        doc.add_paragraph(f"Body paragraph with {MARKERS['body']} inside it.")

        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = f"Table cell with {MARKERS['table']}."

        section = doc.sections[0]
        section.header.paragraphs[0].text = f"Header with {MARKERS['header']}."
        section.footer.paragraphs[0].text = f"Footer with {MARKERS['footer']}."

        doc.save(path)
        self._inject_footnote_endnote_comment_textbox_sdt(path)

    def _inject_footnote_endnote_comment_textbox_sdt(self, path: str) -> None:
        content_types = self._read_zip_entry(path, "[Content_Types].xml")
        doc_rels = self._read_zip_entry(path, "word/_rels/document.xml.rels")

        footnotes_xml = (
            b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            b'<w:footnotes xmlns:w="' + W_NS.encode() + b'">'
            b'<w:footnote w:type="separator" w:id="-1"><w:p/></w:footnote>'
            b'<w:footnote w:type="continuationSeparator" w:id="0"><w:p/></w:footnote>'
            b'<w:footnote w:id="1"><w:p><w:r><w:t>Footnote text '
            + MARKERS["footnote"].encode() + b'</w:t></w:r></w:p></w:footnote>'
            b'</w:footnotes>'
        )
        content_types = self._add_content_type_override(
            content_types, "/word/footnotes.xml",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml",
        )
        doc_rels = self._add_relationship(
            doc_rels, "rIdFootnotes",
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes",
            "footnotes.xml",
        )

        endnotes_xml = (
            b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            b'<w:endnotes xmlns:w="' + W_NS.encode() + b'">'
            b'<w:endnote w:type="separator" w:id="-1"><w:p/></w:endnote>'
            b'<w:endnote w:type="continuationSeparator" w:id="0"><w:p/></w:endnote>'
            b'<w:endnote w:id="1"><w:p><w:r><w:t>Endnote text '
            + MARKERS["endnote"].encode() + b'</w:t></w:r></w:p></w:endnote>'
            b'</w:endnotes>'
        )
        content_types = self._add_content_type_override(
            content_types, "/word/endnotes.xml",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.endnotes+xml",
        )
        doc_rels = self._add_relationship(
            doc_rels, "rIdEndnotes",
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes",
            "endnotes.xml",
        )

        comments_xml = (
            b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            b'<w:comments xmlns:w="' + W_NS.encode() + b'">'
            b'<w:comment w:id="1" w:author="Tester" w:date="2026-01-01T00:00:00Z">'
            b'<w:p><w:r><w:t>Comment text ' + MARKERS["comment"].encode() + b'</w:t></w:r></w:p>'
            b'</w:comment></w:comments>'
        )
        content_types = self._add_content_type_override(
            content_types, "/word/comments.xml",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml",
        )
        doc_rels = self._add_relationship(
            doc_rels, "rIdComments",
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments",
            "comments.xml",
        )

        self._patch_zip(
            path,
            {
                "[Content_Types].xml": content_types,
                "word/_rels/document.xml.rels": doc_rels,
            },
            {
                "word/footnotes.xml": footnotes_xml,
                "word/endnotes.xml": endnotes_xml,
                "word/comments.xml": comments_xml,
            },
        )

        document_xml = self._read_zip_entry(path, "word/document.xml")
        root = ET.fromstring(document_xml)
        body = root.find(_qn(W_NS, "body"))

        # Footnote/endnote references (so the parts are demonstrably "used").
        para_refs = ET.SubElement(body, _qn(W_NS, "p"))
        run_fn = ET.SubElement(para_refs, _qn(W_NS, "r"))
        ET.SubElement(run_fn, _qn(W_NS, "footnoteReference"), {_qn(W_NS, "id"): "1"})
        run_en = ET.SubElement(para_refs, _qn(W_NS, "r"))
        ET.SubElement(run_en, _qn(W_NS, "endnoteReference"), {_qn(W_NS, "id"): "1"})

        # Comment range + reference wrapping a dummy run.
        ET.SubElement(body, _qn(W_NS, "commentRangeStart"), {_qn(W_NS, "id"): "1"})
        para_commented = ET.SubElement(body, _qn(W_NS, "p"))
        run_commented = ET.SubElement(para_commented, _qn(W_NS, "r"))
        t_commented = ET.SubElement(run_commented, _qn(W_NS, "t"))
        t_commented.text = "Commented-upon body text."
        ET.SubElement(body, _qn(W_NS, "commentRangeEnd"), {_qn(W_NS, "id"): "1"})
        para_ref = ET.SubElement(body, _qn(W_NS, "p"))
        run_ref = ET.SubElement(para_ref, _qn(W_NS, "r"))
        ET.SubElement(run_ref, _qn(W_NS, "commentReference"), {_qn(W_NS, "id"): "1"})

        # Legacy VML textbox: w:pict > v:shape > v:textbox > w:txbxContent > w:p
        para_tb = ET.SubElement(body, _qn(W_NS, "p"))
        run_tb = ET.SubElement(para_tb, _qn(W_NS, "r"))
        pict = ET.SubElement(run_tb, _qn(W_NS, "pict"))
        shape = ET.SubElement(pict, _qn(V_NS, "shape"))
        textbox = ET.SubElement(shape, _qn(V_NS, "textbox"))
        txbx_content = ET.SubElement(textbox, _qn(W_NS, "txbxContent"))
        tb_p = ET.SubElement(txbx_content, _qn(W_NS, "p"))
        tb_r = ET.SubElement(tb_p, _qn(W_NS, "r"))
        tb_t = ET.SubElement(tb_r, _qn(W_NS, "t"))
        tb_t.text = f"Textbox text {MARKERS['textbox']}"

        # Block-level content control: w:sdt > w:sdtContent > w:p > w:r > w:t
        # (CT_SdtContentBlock requires w:p/w:tbl/w:sdt -- never a bare w:r --
        # this matches what Word actually emits.)
        sdt = ET.SubElement(body, _qn(W_NS, "sdt"))
        ET.SubElement(sdt, _qn(W_NS, "sdtPr"))
        sdt_content = ET.SubElement(sdt, _qn(W_NS, "sdtContent"))
        sdt_p = ET.SubElement(sdt_content, _qn(W_NS, "p"))
        sdt_r = ET.SubElement(sdt_p, _qn(W_NS, "r"))
        sdt_t = ET.SubElement(sdt_r, _qn(W_NS, "t"))
        sdt_t.text = f"SDT text {MARKERS['sdt']}"

        # Inline content control nested in an ordinary paragraph (the common
        # "plain text content control" pattern): w:p > w:sdt > w:sdtContent
        # > w:r > w:t
        para_inline_sdt = ET.SubElement(body, _qn(W_NS, "p"))
        inline_sdt = ET.SubElement(para_inline_sdt, _qn(W_NS, "sdt"))
        ET.SubElement(inline_sdt, _qn(W_NS, "sdtPr"))
        inline_sdt_content = ET.SubElement(inline_sdt, _qn(W_NS, "sdtContent"))
        inline_sdt_r = ET.SubElement(inline_sdt_content, _qn(W_NS, "r"))
        inline_sdt_t = ET.SubElement(inline_sdt_r, _qn(W_NS, "t"))
        inline_sdt_t.text = f"Inline SDT text {MARKERS['sdt_inline']}"

        updated_doc = ET.tostring(root, encoding="utf-8", xml_declaration=True)
        self._patch_zip(path, {"word/document.xml": updated_doc})

    # ---- pipeline runner ---------------------------------------------

    def _run_pipeline(self, output_name: str, report_name: str, metadata_args: str = "",
                       metadata_preset: str | None = None):
        output_docx = os.path.join(self.work_dir, output_name)
        report_json = os.path.join(self.work_dir, report_name)
        env_keys = (
            "MARCUT_METADATA_PRESET",
            "MARCUT_METADATA_SETTINGS_JSON",
            "MARCUT_METADATA_ARGS",
            "MARCUT_SCRUB_REPORT_PATH",
        )
        previous = {k: os.environ.get(k) for k in env_keys}
        try:
            if metadata_preset is None:
                os.environ.pop("MARCUT_METADATA_PRESET", None)
            else:
                os.environ["MARCUT_METADATA_PRESET"] = metadata_preset
            os.environ.pop("MARCUT_METADATA_SETTINGS_JSON", None)
            os.environ["MARCUT_METADATA_ARGS"] = metadata_args
            os.environ.pop("MARCUT_SCRUB_REPORT_PATH", None)
            code, _timings = pipeline.run_redaction(
                input_path=self.input_docx,
                output_path=output_docx,
                report_path=report_json,
                mode="rules",
                model_id="mock",
                chunk_tokens=200,
                overlap=20,
                temperature=0.1,
                seed=123,
                debug=False,
                backend="mock",
            )
        finally:
            for k, v in previous.items():
                if v is None:
                    os.environ.pop(k, None)
                else:
                    os.environ[k] = v
        self.assertEqual(code, 0, "run_redaction did not report success")
        return output_docx, report_json

    @staticmethod
    def _marker_fully_redacted(zf: zipfile.ZipFile, part_name: str, marker_local: str) -> bool:
        """True only if marker_local shows up exclusively inside w:delText
        (never as live w:t), i.e. it was actually wrapped in a deletion."""
        root = LET.fromstring(zf.read(part_name))
        found_del_text = False
        for el in root.iter():
            local = LET.QName(el.tag).localname if isinstance(el.tag, str) else None
            if local not in ("t", "delText") or not el.text or marker_local not in el.text:
                continue
            if local == "delText":
                found_del_text = True
            else:
                return False  # still live, unredacted plain text
        return found_del_text

    # ---- tests ---------------------------------------------------------

    def test_body_table_header_footer_footnote_endnote_textbox_sdt_all_redacted(self):
        """Every part that already gets scanned must also get written back
        with real track-changes -- not merely indexed at build time.

        Uses the "none" metadata preset to isolate the redaction/track-changes
        layer from the separate metadata-hardening feature, which -- by
        default and by design -- wholesale-removes headers/footers
        (clean_headers_footers) the same way it wholesale-removes comments.
        That default-removal path is covered by
        test_default_settings_remove_comments_entirely_no_leak below; this
        test instead answers "if a caller keeps their headers/footers (e.g.
        --no-clean-headers-footers, to preserve a required confidentiality
        legend), does the PII inside them still get redacted?"
        """
        output_docx, _report_json = self._run_pipeline(
            "output.docx", "report.json", metadata_preset="none",
        )

        redacted_parts = {
            "body": "word/document.xml",
            "table": "word/document.xml",
            "header": "word/header1.xml",
            "footer": "word/footer1.xml",
            "footnote": "word/footnotes.xml",
            "endnote": "word/endnotes.xml",
            "textbox": "word/document.xml",
            "sdt": "word/document.xml",
            "sdt_inline": "word/document.xml",
        }
        with zipfile.ZipFile(output_docx) as zf:
            names = set(zf.namelist())
            for label, part_name in redacted_parts.items():
                self.assertIn(part_name, names, f"{label}: part {part_name} missing from output")
                marker_local = MARKERS[label].split("@", 1)[0]
                self.assertTrue(
                    self._marker_fully_redacted(zf, part_name, marker_local),
                    f"{label}: marker not fully wrapped in w:delText inside {part_name}",
                )

    def test_default_settings_remove_comments_entirely_no_leak(self):
        """Out of the box, review comments AND headers/footers are
        wholesale-removed by the metadata-hardening defaults
        (clean_review_comments_*, clean_headers_footers), so no PII can leak
        through either unscanned part even though neither's text is ever
        scanned by the redaction engine itself."""
        output_docx, report_json = self._run_pipeline("output_default.docx", "report_default.json")

        with zipfile.ZipFile(output_docx) as zf:
            names = zf.namelist()
            self.assertNotIn(
                "word/comments.xml", names,
                "Default settings should remove the comments part entirely",
            )
            self.assertFalse(
                any(n.startswith("word/header") or n.startswith("word/footer") for n in names),
                "Default settings should remove header/footer parts entirely",
            )

        with open(report_json) as f:
            report = json.load(f)
        codes = {w.get("code") for w in report.get("warnings", [])}
        self.assertNotIn(
            "REVIEW_COMMENTS_NOT_SCANNED", codes,
            "No warning expected when comments are fully removed by default",
        )

    def test_retained_visible_comment_is_unredacted_but_flagged(self):
        """If a caller explicitly retains visible comments, the comment's
        PII is neither scanned nor redacted (documented limitation) -- but
        that gap must be disclosed via a report warning, not left silent."""
        output_docx, report_json = self._run_pipeline(
            "output_keepcomments.docx",
            "report_keepcomments.json",
            metadata_args="--no-clean-review-comments-visible",
        )

        with zipfile.ZipFile(output_docx) as zf:
            names = zf.namelist()
            self.assertIn("word/comments.xml", names)
            raw = zf.read("word/comments.xml")
            self.assertIn(MARKERS["comment"].encode(), raw)

        with open(report_json) as f:
            report = json.load(f)
        codes = {w.get("code") for w in report.get("warnings", [])}
        self.assertIn(
            "REVIEW_COMMENTS_NOT_SCANNED", codes,
            "Retaining comments must surface a report warning instead of failing silently",
        )


if __name__ == "__main__":
    unittest.main()

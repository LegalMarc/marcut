"""
Synthetic labeled PII corpus for the precision/recall evaluation harness (A1).

Builds a small DOCX fixture *at run time* (never committed as a binary -- CI's
hygiene job forbids checked-in .docx files) covering each core entity type the
rules engine can detect -- EMAIL, PHONE, SSN, CARD (credit card), MONEY, DATE,
ORG, LOC (address), NAME -- spread across the document body, a table cell, the
header, and the footer.

The gold labels live in `labels.json` alongside this module, not as Python
literals, so they can be reviewed/updated independently of the generator code.
Keep `labels.json` and `build_corpus_docx()` below in sync: every entity text
in the JSON file must appear verbatim in the document this module builds.
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Dict, List

try:
    from docx import Document
    DOCX_AVAILABLE = True
except Exception:  # pragma: no cover - exercised only in stripped-down environments
    DOCX_AVAILABLE = False

_LABELS_PATH = Path(__file__).resolve().parent / "labels.json"


def load_expected_entities() -> List[Dict[str, str]]:
    """Load the gold-label entity list from labels.json."""
    data = json.loads(_LABELS_PATH.read_text(encoding="utf-8"))
    return data["entities"]


def build_corpus_docx(path: str) -> None:
    """Write the synthetic labeled corpus to `path` as a .docx file.

    Every PII value planted here must match an entry in labels.json exactly
    (see that file's `text` fields).
    """
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx is required to build the PII eval corpus")

    doc = Document()

    # --- Body ---
    doc.add_paragraph(
        "This engagement letter is issued to confirm the terms discussed below."
    )
    doc.add_paragraph(
        "Please direct all correspondence to Jane Doe at jane.doe@example.com "
        "or by phone at (415) 555-0182."
    )
    doc.add_paragraph(
        "The client's SSN on record is 123-45-6789, and the retainer card "
        "charged ends in 4111-1111-1111-1111."
    )
    doc.add_paragraph(
        "An initial retainer of $12,500.00 is due no later than January 15, 2024."
    )
    doc.add_paragraph(
        "This Agreement is entered into between Acme Redwood Corp. and the "
        "undersigned client."
    )
    doc.add_paragraph(
        "The subject property is located at 742 Evergreen Terrace, Springfield, "
        "IL 62704."
    )
    doc.add_paragraph("Name:   Sam Jacobs")

    # --- Table ---
    # Note: cell 0 is deliberately filler (no PII) rather than another NAME
    # signature line. A signature line placed immediately before an ORG-suffix
    # line (nothing but capitalized words between them) can get absorbed into
    # one bogus COMPANY_SUFFIX match, since that regex's inter-token separator
    # (`\s+`) crosses paragraph/cell newlines -- a pre-existing rules.py
    # accuracy quirk, out of scope for this harness (see docs/backlog
    # item A6). Keeping filler here avoids coupling this test to that bug.
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Row 1 details:"
    table.rows[0].cells[1].text = "Vertex Analytics Group LLC"

    # --- Header / footer ---
    section = doc.sections[0]
    section.header.paragraphs[0].text = "Name:   Jordan Casey"
    section.header.add_paragraph("Header contact: jordan.casey@example.com")
    section.footer.paragraphs[0].text = "Name:   Taylor Brooks"

    doc.save(path)

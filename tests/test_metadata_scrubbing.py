#!/usr/bin/env python3
"""
Metadata scrubbing regression tests.
Focuses on:
- Settings parsing and CLI flags
- Report shape (groups + before/after values)
- ProofState validity after scrubbing
"""

import json
import os
import tempfile
import unittest
import zipfile
from dataclasses import fields
from pathlib import Path
from xml.etree import ElementTree as ET

try:
    from docx import Document
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

try:
    from marcut.docx_io import MetadataCleaningSettings
    from marcut import pipeline
    IMPORTS_SUCCESS = True
except Exception:
    IMPORTS_SUCCESS = False


class TestMetadataCleaningSettings(unittest.TestCase):
    """Test MetadataCleaningSettings dataclass and CLI parsing."""

    @unittest.skipUnless(IMPORTS_SUCCESS, "marcut.docx_io not available")
    def test_default_settings(self):
        settings = MetadataCleaningSettings()
        self.assertTrue(settings.clean_company)
        self.assertTrue(settings.clean_author)
        self.assertFalse(settings.clean_created_date)
        self.assertFalse(settings.clean_modified_date)
        self.assertTrue(settings.clean_language_settings)
        self.assertTrue(settings.clean_ole_sources)

    @unittest.skipUnless(IMPORTS_SUCCESS, "marcut.docx_io not available")
    def test_from_cli_args(self):
        args = [
            "--no-clean-company",
            "--no-clean-language-settings",
            "--no-clean-ole-sources",
        ]
        settings = MetadataCleaningSettings.from_cli_args(args)
        self.assertFalse(settings.clean_company)
        self.assertFalse(settings.clean_language_settings)
        self.assertFalse(settings.clean_ole_sources)

    @unittest.skipUnless(IMPORTS_SUCCESS, "marcut.docx_io not available")
    def test_to_cli_args_none_preset(self):
        settings = MetadataCleaningSettings()
        for f in fields(settings):
            setattr(settings, f.name, False)
        args = settings.to_cli_args()
        self.assertIn("--preset-none", args)


class TestMetadataScrubReport(unittest.TestCase):
    """Integration tests that require python-docx."""

    @staticmethod
    def _read_zip_entry(path: str, name: str) -> bytes:
        with zipfile.ZipFile(path) as zf:
            return zf.read(name)

    @staticmethod
    def _patch_zip(path: str, updates: dict, new_entries: dict | None = None) -> None:
        new_entries = new_entries or {}
        temp_path = path + ".tmp"
        with zipfile.ZipFile(path, "r") as zin:
            existing = {item.filename for item in zin.infolist()}
            with zipfile.ZipFile(temp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    name = item.filename
                    data = updates.get(name, zin.read(name))
                    zout.writestr(item, data)
                for name, data in new_entries.items():
                    if name not in existing:
                        zout.writestr(name, data)
        os.replace(temp_path, path)

    @staticmethod
    def _add_override(content_types: bytes, part_name: str, content_type: str) -> bytes:
        root = ET.fromstring(content_types)
        ns = root.tag.split("}")[0].strip("{")
        for override in root.findall(f"{{{ns}}}Override"):
            if override.get("PartName") == part_name:
                return content_types
        ET.SubElement(root, f"{{{ns}}}Override", {
            "PartName": part_name,
            "ContentType": content_type,
        })
        return ET.tostring(root, encoding="utf-8", xml_declaration=True)

    @staticmethod
    def _add_relationship(rels_xml: bytes, rel_id: str, rel_type: str, target: str) -> bytes:
        root = ET.fromstring(rels_xml)
        ns = root.tag.split("}")[0].strip("{")
        for rel in root.findall(f"{{{ns}}}Relationship"):
            if rel.get("Id") == rel_id:
                return rels_xml
        ET.SubElement(root, f"{{{ns}}}Relationship", {
            "Id": rel_id,
            "Type": rel_type,
            "Target": target,
        })
        return ET.tostring(root, encoding="utf-8", xml_declaration=True)

    def _make_docx_with_header(self, path: str) -> None:
        doc = Document()
        doc.add_paragraph("Metadata scrub test document.")
        header = doc.sections[0].header
        header.paragraphs[0].text = "Header content"
        doc.save(path)

    def _inject_mail_merge(self, path: str) -> None:
        document_xml = self._read_zip_entry(path, "word/document.xml")
        root = ET.fromstring(document_xml)
        w_ns = root.tag.split("}")[0].strip("{")
        body = root.find(f"{{{w_ns}}}body")
        para = ET.SubElement(body, f"{{{w_ns}}}p")
        fld = ET.SubElement(para, f"{{{w_ns}}}fldSimple")
        fld.set(f"{{{w_ns}}}instr", " MERGEFIELD FirstName \\* MERGEFORMAT ")
        run = ET.SubElement(fld, f"{{{w_ns}}}r")
        text = ET.SubElement(run, f"{{{w_ns}}}t")
        text.text = "FirstName"
        updated_doc = ET.tostring(root, encoding="utf-8", xml_declaration=True)

        settings_xml = self._read_zip_entry(path, "word/settings.xml")
        settings_root = ET.fromstring(settings_xml)
        w_ns = settings_root.tag.split("}")[0].strip("{")
        mail_merge = ET.SubElement(settings_root, f"{{{w_ns}}}mailMerge")
        ET.SubElement(mail_merge, f"{{{w_ns}}}dataSource")
        ET.SubElement(mail_merge, f"{{{w_ns}}}headerSource")
        updated_settings = ET.tostring(settings_root, encoding="utf-8", xml_declaration=True)

        self._patch_zip(path, {
            "word/document.xml": updated_doc,
            "word/settings.xml": updated_settings,
        })

    def _inject_hidden_text(self, path: str) -> None:
        document_xml = self._read_zip_entry(path, "word/document.xml")
        root = ET.fromstring(document_xml)
        w_ns = root.tag.split("}")[0].strip("{")
        body = root.find(f"{{{w_ns}}}body")
        para = ET.SubElement(body, f"{{{w_ns}}}p")
        run = ET.SubElement(para, f"{{{w_ns}}}r")
        rpr = ET.SubElement(run, f"{{{w_ns}}}rPr")
        ET.SubElement(rpr, f"{{{w_ns}}}vanish")
        text = ET.SubElement(run, f"{{{w_ns}}}t")
        text.text = "Hidden content"
        updated_doc = ET.tostring(root, encoding="utf-8", xml_declaration=True)
        self._patch_zip(path, {"word/document.xml": updated_doc})

    def _inject_data_binding(self, path: str) -> None:
        document_xml = self._read_zip_entry(path, "word/document.xml")
        root = ET.fromstring(document_xml)
        w_ns = root.tag.split("}")[0].strip("{")
        body = root.find(f"{{{w_ns}}}body")
        sdt = ET.SubElement(body, f"{{{w_ns}}}sdt")
        sdt_pr = ET.SubElement(sdt, f"{{{w_ns}}}sdtPr")
        ET.SubElement(sdt_pr, f"{{{w_ns}}}dataBinding")
        sdt_content = ET.SubElement(sdt, f"{{{w_ns}}}sdtContent")
        run = ET.SubElement(sdt_content, f"{{{w_ns}}}r")
        text = ET.SubElement(run, f"{{{w_ns}}}t")
        text.text = "Bound content"
        updated_doc = ET.tostring(root, encoding="utf-8", xml_declaration=True)
        self._patch_zip(path, {"word/document.xml": updated_doc})

    def _inject_invisible_shape(self, path: str) -> None:
        document_xml = self._read_zip_entry(path, "word/document.xml")
        root = ET.fromstring(document_xml)
        w_ns = root.tag.split("}")[0].strip("{")
        body = root.find(f"{{{w_ns}}}body")
        para = ET.SubElement(body, f"{{{w_ns}}}p")
        run = ET.SubElement(para, f"{{{w_ns}}}r")
        pict = ET.SubElement(run, f"{{{w_ns}}}pict")
        v_ns = "urn:schemas-microsoft-com:vml"
        shape = ET.SubElement(pict, f"{{{v_ns}}}shape")
        shape.set("style", "visibility:hidden")
        updated_doc = ET.tostring(root, encoding="utf-8", xml_declaration=True)
        self._patch_zip(path, {"word/document.xml": updated_doc})

    def _inject_hyperlink_base(self, path: str) -> None:
        app_xml = self._read_zip_entry(path, "docProps/app.xml")
        root = ET.fromstring(app_xml)
        ep_ns = root.tag.split("}")[0].strip("{")
        hb = ET.SubElement(root, f"{{{ep_ns}}}HyperlinkBase")
        hb.text = "https://example.com/base"
        updated_app = ET.tostring(root, encoding="utf-8", xml_declaration=True)
        self._patch_zip(path, {"docProps/app.xml": updated_app})

    def _inject_custom_xml(self, path: str) -> None:
        content_types = self._read_zip_entry(path, "[Content_Types].xml")
        content_types = self._add_override(
            content_types,
            "/docProps/custom.xml",
            "application/vnd.openxmlformats-officedocument.custom-properties+xml",
        )
        content_types = self._add_override(
            content_types,
            "/customXml/item1.xml",
            "application/xml",
        )
        custom_props = (
            b'<?xml version="1.0" encoding="UTF-8"?>'
            b'<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/custom-properties" '
            b'xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">'
            b'<property fmtid="{D5CDD505-2E9C-101B-9397-08002B2CF9AE}" pid="2" name="ClientId">'
            b'<vt:lpwstr>ABC-123</vt:lpwstr></property></Properties>'
        )
        custom_xml = b'<?xml version="1.0" encoding="UTF-8"?><foo xmlns="http://example.com/custom">bar</foo>'
        rels_xml = self._read_zip_entry(path, "_rels/.rels")
        rels_xml = self._add_relationship(
            rels_xml,
            "rIdCustomProps",
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/custom-properties",
            "docProps/custom.xml",
        )
        doc_rels = self._read_zip_entry(path, "word/_rels/document.xml.rels")
        doc_rels = self._add_relationship(
            doc_rels,
            "rIdCustomXml",
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/customXml",
            "../customXml/item1.xml",
        )
        self._patch_zip(
            path,
            {
                "[Content_Types].xml": content_types,
                "_rels/.rels": rels_xml,
                "word/_rels/document.xml.rels": doc_rels,
            },
            {
                "docProps/custom.xml": custom_props,
                "customXml/item1.xml": custom_xml,
            },
        )

    @staticmethod
    def _find_part(path: str, prefix: str) -> str | None:
        with zipfile.ZipFile(path) as zf:
            for name in zf.namelist():
                if name.startswith(prefix):
                    return name
        return None

    def _inject_watermark(self, path: str) -> None:
        header_part = self._find_part(path, "word/header")
        if not header_part:
            raise AssertionError("header part not found for watermark injection")
        header_xml = self._read_zip_entry(path, header_part)
        root = ET.fromstring(header_xml)
        w_ns = root.tag.split("}")[0].strip("{")
        para = root.find(f"{{{w_ns}}}p")
        if para is None:
            para = ET.SubElement(root, f"{{{w_ns}}}p")
        run = ET.SubElement(para, f"{{{w_ns}}}r")
        pict = ET.SubElement(run, f"{{{w_ns}}}pict")
        v_ns = "urn:schemas-microsoft-com:vml"
        shape = ET.SubElement(pict, f"{{{v_ns}}}shape")
        shape.set("id", "PowerPlusWaterMarkObject")
        shape.set("style", "mso-position-horizontal:center; mso-position-vertical:center; z-index:-1")
        updated_header = ET.tostring(root, encoding="utf-8", xml_declaration=True)
        self._patch_zip(path, {header_part: updated_header})

    def _inject_document_versions(self, path: str) -> None:
        content_types = self._read_zip_entry(path, "[Content_Types].xml")
        content_types = self._add_override(content_types, "/word/versions/versions.xml", "application/xml")
        doc_rels = self._read_zip_entry(path, "word/_rels/document.xml.rels")
        doc_rels = self._add_relationship(
            doc_rels,
            "rIdVersions",
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/versions",
            "versions/versions.xml",
        )
        versions_xml = b'<?xml version="1.0" encoding="UTF-8"?><versions xmlns="http://example.com/versions" />'
        self._patch_zip(
            path,
            {
                "[Content_Types].xml": content_types,
                "word/_rels/document.xml.rels": doc_rels,
            },
            {"word/versions/versions.xml": versions_xml},
        )

    def _inject_ink_annotations(self, path: str) -> None:
        content_types = self._read_zip_entry(path, "[Content_Types].xml")
        content_types = self._add_override(content_types, "/word/ink/ink1.xml", "application/xml")
        doc_rels = self._read_zip_entry(path, "word/_rels/document.xml.rels")
        doc_rels = self._add_relationship(
            doc_rels,
            "rIdInk",
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/ink",
            "ink/ink1.xml",
        )
        ink_xml = b'<?xml version="1.0" encoding="UTF-8"?><ink xmlns="http://example.com/ink" />'
        self._patch_zip(
            path,
            {
                "[Content_Types].xml": content_types,
                "word/_rels/document.xml.rels": doc_rels,
            },
            {"word/ink/ink1.xml": ink_xml},
        )

    def _copy_docx(self, source: str, name: str) -> str:
        import shutil

        path = os.path.join(self.temp_dir, name)
        shutil.copy2(source, path)
        return path

    def _run_scrub(self, input_path: str, output_name: str, metadata_args: str = "") -> tuple[str, dict]:
        output_docx = os.path.join(self.temp_dir, output_name)
        prev_args = os.environ.get("MARCUT_METADATA_ARGS")
        os.environ["MARCUT_METADATA_ARGS"] = metadata_args
        try:
            success, error, report = pipeline.scrub_metadata_only(
                input_path=input_path,
                output_path=output_docx,
                debug=False,
            )
        finally:
            if prev_args is None:
                os.environ.pop("MARCUT_METADATA_ARGS", None)
            else:
                os.environ["MARCUT_METADATA_ARGS"] = prev_args
        self.assertTrue(success, msg=error or "scrub failed")
        return output_docx, report

    @staticmethod
    def _extract_text(xml_bytes: bytes) -> str:
        root = ET.fromstring(xml_bytes)
        texts = []
        for el in root.iter():
            if el.tag.endswith("}t") and el.text:
                texts.append(el.text)
        return "".join(texts)

    @classmethod
    def setUpClass(cls):
        cls.temp_dir = tempfile.mkdtemp(prefix="marcut_metadata_")
        cls.input_docx = None
        if not (DOCX_AVAILABLE and IMPORTS_SUCCESS):
            return

        doc = Document()
        doc.add_paragraph("Metadata scrub test document.")
        doc.core_properties.author = "Test Author"
        doc.core_properties.title = "Test Title"
        cls.input_docx = os.path.join(cls.temp_dir, "input.docx")
        doc.save(cls.input_docx)

    @classmethod
    def tearDownClass(cls):
        if os.path.exists(cls.temp_dir):
            import shutil
            shutil.rmtree(cls.temp_dir)

    @unittest.skipUnless(DOCX_AVAILABLE and IMPORTS_SUCCESS, "python-docx or marcut not available")
    def test_scrub_report_groups(self):
        output_docx, report = self._run_scrub(self.input_docx, "output.docx")
        self.assertIn("groups", report)
        self.assertIn("summary", report)

        expected_groups = {
            "App Properties",
            "Core Properties",
            "Custom Properties",
            "Document Structure",
            "Embedded Content",
            "Advanced Hardening",
        }
        self.assertEqual(set(report["groups"].keys()), expected_groups)

        for group_name, items in report["groups"].items():
            self.assertIsInstance(items, list)
            self.assertGreater(len(items), 0, msg=f"{group_name} empty")
            for item in items:
                self.assertIn("field", item)
                self.assertIn("before", item)
                self.assertIn("after", item)
                self.assertIn("status", item)

    @unittest.skipUnless(DOCX_AVAILABLE and IMPORTS_SUCCESS, "python-docx or marcut not available")
    def test_scrub_report_includes_new_fields(self):
        output_docx, report = self._run_scrub(self.input_docx, "report_fields.docx")
        self.assertTrue(os.path.exists(output_docx))

        app_fields = {item["field"] for item in report["groups"]["App Properties"]}
        self.assertIn("Hyperlink Base", app_fields)

        custom_fields = {item["field"] for item in report["groups"]["Custom Properties"]}
        self.assertIn("Custom Properties & Custom XML", custom_fields)

        doc_fields = {item["field"] for item in report["groups"]["Document Structure"]}
        for name in [
            "Mail Merge Data",
            "Data Bindings",
            "Document Versions",
            "Ink Annotations",
            "Hidden Text",
            "Invisible Objects",
            "Headers & Footers",
            "Watermarks",
        ]:
            self.assertIn(name, doc_fields)

    @unittest.skipUnless(DOCX_AVAILABLE and IMPORTS_SUCCESS, "python-docx or marcut not available")
    def test_mail_merge_cleanup(self):
        input_docx = self._copy_docx(self.input_docx, "mail_merge_input.docx")
        self._inject_mail_merge(input_docx)
        output_docx, _report = self._run_scrub(input_docx, "mail_merge_output.docx")

        settings_xml = self._read_zip_entry(output_docx, "word/settings.xml")
        self.assertNotIn(b"mailMerge", settings_xml)

        doc_xml = self._read_zip_entry(output_docx, "word/document.xml")
        text = self._extract_text(doc_xml)
        self.assertIn("<<FirstName>>", text)
        self.assertNotIn("MERGEFIELD", text)
        self.assertNotIn(b"MERGEFIELD", doc_xml)

    @unittest.skipUnless(DOCX_AVAILABLE and IMPORTS_SUCCESS, "python-docx or marcut not available")
    def test_data_bindings_removed(self):
        input_docx = self._copy_docx(self.input_docx, "data_bindings_input.docx")
        self._inject_data_binding(input_docx)
        output_docx, _report = self._run_scrub(input_docx, "data_bindings_output.docx")

        doc_xml = self._read_zip_entry(output_docx, "word/document.xml")
        root = ET.fromstring(doc_xml)
        self.assertEqual(sum(1 for el in root.iter() if el.tag.endswith("}dataBinding")), 0)

    @unittest.skipUnless(DOCX_AVAILABLE and IMPORTS_SUCCESS, "python-docx or marcut not available")
    def test_hidden_text_removed(self):
        input_docx = self._copy_docx(self.input_docx, "hidden_text_input.docx")
        self._inject_hidden_text(input_docx)
        output_docx, _report = self._run_scrub(input_docx, "hidden_text_output.docx")

        doc_xml = self._read_zip_entry(output_docx, "word/document.xml")
        root = ET.fromstring(doc_xml)
        self.assertEqual(sum(1 for el in root.iter() if el.tag.endswith("}vanish")), 0)
        self.assertNotIn("Hidden content", self._extract_text(doc_xml))

    @unittest.skipUnless(DOCX_AVAILABLE and IMPORTS_SUCCESS, "python-docx or marcut not available")
    def test_invisible_objects_removed(self):
        input_docx = self._copy_docx(self.input_docx, "invisible_input.docx")
        self._inject_invisible_shape(input_docx)
        output_docx, _report = self._run_scrub(input_docx, "invisible_output.docx")

        doc_xml = self._read_zip_entry(output_docx, "word/document.xml")
        self.assertNotIn(b"visibility:hidden", doc_xml)
        self.assertNotIn(b"visibility: hidden", doc_xml)

    @unittest.skipUnless(DOCX_AVAILABLE and IMPORTS_SUCCESS, "python-docx or marcut not available")
    def test_hyperlink_base_removed(self):
        input_docx = self._copy_docx(self.input_docx, "hyperlink_base_input.docx")
        self._inject_hyperlink_base(input_docx)
        output_docx, _report = self._run_scrub(input_docx, "hyperlink_base_output.docx")

        app_xml = self._read_zip_entry(output_docx, "docProps/app.xml")
        root = ET.fromstring(app_xml)
        self.assertEqual(sum(1 for el in root.iter() if el.tag.endswith("}HyperlinkBase")), 0)

    @unittest.skipUnless(DOCX_AVAILABLE and IMPORTS_SUCCESS, "python-docx or marcut not available")
    def test_custom_xml_report_details(self):
        input_docx = self._copy_docx(self.input_docx, "custom_xml_input.docx")
        self._inject_custom_xml(input_docx)
        output_docx, report = self._run_scrub(input_docx, "custom_xml_output.docx")

        custom_group = report["groups"]["Custom Properties"]
        entry = next(item for item in custom_group if item["field"] == "Custom Properties & Custom XML")
        before = entry["before"]
        self.assertIn("custom_property_names", before)
        self.assertIn("ClientId", before["custom_property_names"])
        self.assertTrue(any(part.get("part") == "/customXml/item1.xml" for part in before["custom_xml_parts"]))
        self.assertGreaterEqual(before["custom_xml_rel_count"], 1)

        with zipfile.ZipFile(output_docx) as zf:
            names = set(zf.namelist())
        self.assertNotIn("docProps/custom.xml", names)
        self.assertNotIn("customXml/item1.xml", names)

    @unittest.skipUnless(DOCX_AVAILABLE and IMPORTS_SUCCESS, "python-docx or marcut not available")
    def test_document_versions_removed(self):
        input_docx = self._copy_docx(self.input_docx, "versions_input.docx")
        self._inject_document_versions(input_docx)
        output_docx, report = self._run_scrub(input_docx, "versions_output.docx")

        doc_group = report["groups"]["Document Structure"]
        versions_entry = next(item for item in doc_group if item["field"] == "Document Versions")
        self.assertEqual(versions_entry["before"], "1 parts")

        with zipfile.ZipFile(output_docx) as zf:
            names = set(zf.namelist())
        self.assertNotIn("word/versions/versions.xml", names)

    @unittest.skipUnless(DOCX_AVAILABLE and IMPORTS_SUCCESS, "python-docx or marcut not available")
    def test_ink_annotations_removed(self):
        input_docx = self._copy_docx(self.input_docx, "ink_input.docx")
        self._inject_ink_annotations(input_docx)
        output_docx, report = self._run_scrub(input_docx, "ink_output.docx")

        doc_group = report["groups"]["Document Structure"]
        ink_entry = next(item for item in doc_group if item["field"] == "Ink Annotations")
        self.assertIn("1 ink parts", ink_entry["before"])

        with zipfile.ZipFile(output_docx) as zf:
            names = set(zf.namelist())
        self.assertNotIn("word/ink/ink1.xml", names)

    @unittest.skipUnless(DOCX_AVAILABLE and IMPORTS_SUCCESS, "python-docx or marcut not available")
    def test_headers_footers_removed(self):
        input_docx = os.path.join(self.temp_dir, "headers_input.docx")
        self._make_docx_with_header(input_docx)
        output_docx, _report = self._run_scrub(input_docx, "headers_output.docx")

        with zipfile.ZipFile(output_docx) as zf:
            names = set(zf.namelist())
        self.assertFalse(any(name.startswith("word/header") for name in names))
        self.assertFalse(any(name.startswith("word/footer") for name in names))

        doc_xml = self._read_zip_entry(output_docx, "word/document.xml")
        root = ET.fromstring(doc_xml)
        self.assertEqual(sum(1 for el in root.iter() if el.tag.endswith("}headerReference")), 0)
        self.assertEqual(sum(1 for el in root.iter() if el.tag.endswith("}footerReference")), 0)

    @unittest.skipUnless(DOCX_AVAILABLE and IMPORTS_SUCCESS, "python-docx or marcut not available")
    def test_watermarks_removed(self):
        input_docx = os.path.join(self.temp_dir, "watermark_input.docx")
        self._make_docx_with_header(input_docx)
        self._inject_watermark(input_docx)
        output_docx, _report = self._run_scrub(
            input_docx,
            "watermark_output.docx",
            metadata_args="--no-clean-headers-footers",
        )

        header_part = self._find_part(output_docx, "word/header")
        self.assertIsNotNone(header_part)
        header_xml = self._read_zip_entry(output_docx, header_part)
        self.assertNotIn(b"WaterMark", header_xml)
        self.assertNotIn(b"watermark", header_xml.lower())

    @unittest.skipUnless(DOCX_AVAILABLE and IMPORTS_SUCCESS, "python-docx or marcut not available")
    def test_proof_state_valid(self):
        output_docx, _report = self._run_scrub(self.input_docx, "proofstate.docx")

        with zipfile.ZipFile(output_docx) as zf:
            data = zf.read("word/settings.xml")
        root = ET.fromstring(data)
        proof_states = [el for el in root.iter() if el.tag.endswith("}proofState")]
        self.assertEqual(len(proof_states), 1)
        attrs = proof_states[0].attrib
        for key in attrs.keys():
            self.assertFalse(key.endswith("}clean"), msg="invalid proofState attribute")
        self.assertTrue(
            any(key.endswith("}spelling") for key in attrs.keys()) or
            any(key.endswith("}grammar") for key in attrs.keys())
        )

    @unittest.skipUnless(DOCX_AVAILABLE and IMPORTS_SUCCESS, "python-docx or marcut not available")
    def test_redaction_writes_scrub_report(self):
        output_docx = os.path.join(self.temp_dir, "redaction.docx")
        audit_report = os.path.join(self.temp_dir, "redaction_report.json")
        scrub_report = os.path.join(self.temp_dir, "scrub_report.json")
        prev_args = os.environ.get("MARCUT_METADATA_ARGS")
        prev_scrub = os.environ.get("MARCUT_SCRUB_REPORT_PATH")
        os.environ["MARCUT_METADATA_ARGS"] = ""
        os.environ["MARCUT_SCRUB_REPORT_PATH"] = scrub_report
        try:
            code, _timings = pipeline.run_redaction(
                input_path=self.input_docx,
                output_path=output_docx,
                report_path=audit_report,
                mode="rules",
                model_id="mock",
                chunk_tokens=200,
                overlap=20,
                temperature=0.1,
                seed=123,
                debug=False,
                backend="mock",
            )
        finally:
            if prev_args is None:
                os.environ.pop("MARCUT_METADATA_ARGS", None)
            else:
                os.environ["MARCUT_METADATA_ARGS"] = prev_args
            if prev_scrub is None:
                os.environ.pop("MARCUT_SCRUB_REPORT_PATH", None)
            else:
                os.environ["MARCUT_SCRUB_REPORT_PATH"] = prev_scrub

        self.assertEqual(code, 0)
        self.assertTrue(os.path.exists(output_docx))
        self.assertTrue(os.path.exists(audit_report))
        self.assertTrue(os.path.exists(scrub_report))
        with open(scrub_report, "r", encoding="utf-8") as fh:
            report = json.load(fh)
        self.assertIn("groups", report)
        self.assertIn("summary", report)


if __name__ == "__main__":
    unittest.main()

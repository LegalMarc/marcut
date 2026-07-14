"""Synthetic large-DOCX production-path performance gates."""

import io
import json
import os
import resource
import struct
import subprocess
import sys
import textwrap
import time
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

import pytest

try:
    from docx import Document
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

try:
    from marcut import pipeline
    IMPORTS_SUCCESS = True
except Exception:
    IMPORTS_SUCCESS = False

try:
    from marcut.docx_io import DocxMap, MetadataCleaningSettings
    DOCX_IO_IMPORTS_SUCCESS = True
except Exception:
    DOCX_IO_IMPORTS_SUCCESS = False


def _add_override(content_types: bytes, part_name: str, content_type: str) -> bytes:
    root = ET.fromstring(content_types)
    ns = root.tag.split("}")[0].strip("{")
    for el in root.findall(f"{{{ns}}}Override"):
        if el.get("PartName") == part_name:
            return content_types
    ET.SubElement(root, f"{{{ns}}}Override", {"PartName": part_name, "ContentType": content_type})
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _add_relationship(rels_xml: bytes, rid: str, reltype: str, target: str) -> bytes:
    root = ET.fromstring(rels_xml)
    ns = root.tag.split("}")[0].strip("{")
    for rel in root.findall(f"{{{ns}}}Relationship"):
        if rel.get("Id") == rid:
            return rels_xml
    ET.SubElement(root, f"{{{ns}}}Relationship", {"Id": rid, "Type": reltype, "Target": target})
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _patch_zip(path: Path, updates: dict[str, bytes], new_entries: dict[str, bytes]) -> None:
    temp_path = path.with_suffix(path.suffix + ".tmp")
    with zipfile.ZipFile(path, "r") as zin:
        with zipfile.ZipFile(temp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            existing = set()
            for item in zin.infolist():
                existing.add(item.filename)
                zout.writestr(item, updates.get(item.filename, zin.read(item.filename)))
            for name, data in new_entries.items():
                if name not in existing:
                    zout.writestr(name, data)
    temp_path.replace(path)


def _inject_comment_metadata(path: Path) -> None:
    with zipfile.ZipFile(path) as zf:
        content_types = zf.read("[Content_Types].xml")
        document_rels = zf.read("word/_rels/document.xml.rels")

    content_types = _add_override(
        content_types,
        "/word/comments.xml",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml",
    )
    document_rels = _add_relationship(
        document_rels,
        "rIdSyntheticComments",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments",
        "comments.xml",
    )
    comments_xml = (
        b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        b'<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        b'<w:comment w:id="0" w:author="Performance Reviewer" w:date="2026-05-13T12:00:00Z">'
        b"<w:p><w:r><w:t>Synthetic review comment with personal data: Jordan Example.</w:t></w:r></w:p>"
        b"</w:comment></w:comments>"
    )
    _patch_zip(
        path,
        {
            "[Content_Types].xml": content_types,
            "word/_rels/document.xml.rels": document_rels,
        },
        {"word/comments.xml": comments_xml},
    )


def _build_synthetic_docx(path: Path) -> None:
    doc = Document()
    props = doc.core_properties
    props.author = "Jordan Example"
    props.last_modified_by = "Taylor Reviewer"
    props.title = "Synthetic Performance Matter"
    props.comments = "Contains generated metadata for the production-path performance gate."

    section = doc.sections[0]
    section.header.paragraphs[0].text = "Header: Jordan Example | Acme Confidential Holdings"
    section.footer.paragraphs[0].text = "Footer contact: jordan.example@example.com | 212-555-0199"

    for idx in range(160):
        doc.add_paragraph(
            f"Paragraph {idx}: Jordan Example at Acme Confidential Holdings LLC uses "
            f"jordan.example{idx}@example.com, account ACCT-{idx:06d}, and phone 212-555-{idx % 10000:04d}."
        )

    table = doc.add_table(rows=36, cols=4)
    for row_idx, row in enumerate(table.rows):
        row.cells[0].text = f"Row {row_idx}"
        row.cells[1].text = f"Acme Confidential Holdings LLC"
        row.cells[2].text = f"Client {row_idx} SSN 123-45-{row_idx % 10000:04d}"
        row.cells[3].text = f"https://example.com/client/{row_idx}"

    doc.save(path)
    _inject_comment_metadata(path)


@pytest.mark.skipif(not (DOCX_AVAILABLE and IMPORTS_SUCCESS), reason="python-docx or marcut not available")
def test_large_docx_rules_pipeline_records_stable_metrics(tmp_path, monkeypatch):
    input_path = tmp_path / "large_input.docx"
    output_path = tmp_path / "large_output.docx"
    report_path = tmp_path / "large_report.json"
    scrub_report_path = tmp_path / "large_scrub_report.json"
    _build_synthetic_docx(input_path)

    monkeypatch.setenv("MARCUT_SCRUB_REPORT_PATH", str(scrub_report_path))
    monkeypatch.setenv("MARCUT_METADATA_REPORT_MAX_STRING_CHARS", "20000")
    monkeypatch.setenv("MARCUT_METADATA_REPORT_MAX_LIST_ITEMS", "200")
    monkeypatch.setenv("MARCUT_CONSISTENCY_MAX_CANDIDATES", "1500")
    start = time.perf_counter()
    code, timings = pipeline.run_redaction(
        str(input_path),
        str(output_path),
        str(report_path),
        mode="rules",
        model_id="rules",
        chunk_tokens=800,
        overlap=80,
        temperature=0.1,
        seed=42,
        debug=False,
        timing=True,
    )
    wall_seconds = time.perf_counter() - start

    assert code == 0
    assert output_path.exists()
    assert report_path.exists()
    assert scrub_report_path.exists()
    assert wall_seconds < float(os.environ.get("MARCUT_LARGE_DOCX_TEST_MAX_SECONDS", "20"))
    assert set(timings) >= {"DOCX_LOAD", "RULES", "POST_PROCESS", "DOCX_SAVE"}

    report = json.loads(report_path.read_text(encoding="utf-8"))
    scrub_report = json.loads(scrub_report_path.read_text(encoding="utf-8"))
    assert len(report.get("spans", [])) >= 200
    assert output_path.stat().st_size > 0
    assert report_path.stat().st_size < 5_000_000
    assert scrub_report_path.stat().st_size < 5_000_000
    assert scrub_report["summary"]["total_cleaned"] >= 1


def _make_fake_jpeg(filler_size: int, width: int = 100, height: int = 100) -> bytes:
    """Build a minimal-but-parseable JFIF JPEG: a real APP0/SOF0 header so
    python-docx can compute an EMU size, a junk APP1 segment so docx_io's
    EXIF-stripping path treats the part as "changed", and a large random-byte
    filler standing in for real (incompressible) photo scan data."""
    out = bytearray(b"\xFF\xD8")  # SOI

    app0 = b"JFIF\x00" + bytes([1, 1, 1]) + struct.pack(">HH", 72, 72) + bytes([0, 0])
    out += b"\xFF\xE0" + struct.pack(">H", len(app0) + 2) + app0

    app1 = b"\x00\x00\x00\x00"
    out += b"\xFF\xE1" + struct.pack(">H", len(app1) + 2) + app1

    sof = bytes([8]) + struct.pack(">HH", height, width) + bytes([1]) + bytes([1, 0x11, 0])
    out += b"\xFF\xC0" + struct.pack(">H", len(sof) + 2) + sof

    out += b"\xFF\xDA" + struct.pack(">H", 2)  # SOS, no header content
    out += os.urandom(filler_size)
    out += b"\xFF\xD9"  # EOI
    return bytes(out)


def _build_media_heavy_docx(path: Path, n_images: int, image_size: int) -> int:
    doc = Document()
    for i in range(20):
        doc.add_paragraph(f"Paragraph {i}: synthetic body text for the memory regression fixture.")
    for _ in range(n_images):
        doc.add_picture(io.BytesIO(_make_fake_jpeg(image_size)))
    doc.save(str(path))
    return path.stat().st_size


@pytest.mark.skipif(not (DOCX_AVAILABLE and DOCX_IO_IMPORTS_SUCCESS), reason="python-docx or marcut.docx_io not available")
def test_metadata_zip_rewrite_peak_memory_bounded(tmp_path):
    """Regression test for issue #52 (C2): docx_io._rewrite_docx_zip must stream the
    archive part-by-part -- read one entry, transform only the parts that need
    scrubbing, write immediately -- instead of buffering every changed part (including
    large embedded media needing EXIF stripping) in memory for the duration of the
    rewrite. Before the streaming fix, peak RSS attributable to this function scaled
    roughly linearly with total document size (~1.2-1.4x observed on 100-300MB
    synthetic fixtures); the streaming rewrite keeps it flat regardless of file size.

    Measured in a fresh subprocess so the assertion reflects only this function's
    footprint, not this test process's own accumulated memory from earlier tests.
    """
    docx_path = tmp_path / "media_heavy.docx"
    # 20 x 5MB EXIF-bearing images: big enough that a per-part-buffering regression
    # would clearly blow past the threshold below, small enough to keep CI fast.
    file_size = _build_media_heavy_docx(docx_path, n_images=20, image_size=5 * 1024 * 1024)

    src_python = str(Path(__file__).resolve().parents[1] / "src" / "python")
    child_script = textwrap.dedent(
        f"""
        import resource, sys
        sys.path.insert(0, {src_python!r})
        from marcut.docx_io import DocxMap, MetadataCleaningSettings

        dm = DocxMap.__new__(DocxMap)
        dm.warnings = []
        settings = MetadataCleaningSettings()  # defaults: clean_image_exif=True, etc.

        rss_before = resource.getrusage(resource.RUSAGE_SELF).ru_maxrss
        dm._rewrite_docx_zip({str(docx_path)!r}, settings)
        rss_after = resource.getrusage(resource.RUSAGE_SELF).ru_maxrss

        unit = 1 if sys.platform == "darwin" else 1024  # macOS reports bytes, Linux KB
        print("RSS_BEFORE", rss_before * unit)
        print("RSS_AFTER", rss_after * unit)
        """
    )
    result = subprocess.run(
        [sys.executable, "-c", child_script],
        capture_output=True,
        text=True,
        timeout=60,
    )
    assert result.returncode == 0, f"child process failed: {result.stderr}"

    values = {}
    for line in result.stdout.splitlines():
        if line.startswith("RSS_"):
            key, val = line.split()
            values[key] = int(val)
    assert "RSS_BEFORE" in values and "RSS_AFTER" in values, result.stdout

    delta = values["RSS_AFTER"] - values["RSS_BEFORE"]
    # Streaming rewrite measured ~0.2-0.5x file size on 60-300MB fixtures; the old
    # two-pass/buffer-everything implementation measured ~1.2-1.6x. 0.9x sits well
    # above the former and well below the latter, with margin for CI-environment noise.
    max_allowed = int(0.9 * file_size)
    assert delta < max_allowed, (
        f"_rewrite_docx_zip peak RSS delta {delta / (1024*1024):.1f} MB exceeded "
        f"{max_allowed / (1024*1024):.1f} MB for a {file_size / (1024*1024):.1f} MB "
        f"document -- possible streaming regression (whole parts being buffered again)"
    )
